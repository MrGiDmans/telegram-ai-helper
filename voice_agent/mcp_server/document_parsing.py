from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


def extract_pages(file_path: str) -> list[dict]:
    """Достать текст из файла постранично (PDF) или целиком (DOCX).

    Определяем формат по расширению файла, а не по mime_type — mime_type
    может быть неточным/отсутствовать, расширение всегда при файле.
    """
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return _extract_docx_pages(file_path)
    return _extract_pdf_pages(file_path)


def _extract_pdf_pages(file_path: str) -> list[dict]:
    pages_data = []
    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                pages_data.append({"page_number": page_number, "text": text})
    return pages_data


def _extract_docx_pages(file_path: str) -> list[dict]:
    # .docx хранит текст как обычный Unicode в XML — в отличие от некоторых PDF
    # (например, экспортированных из LaTeX с неполной ToUnicode-таблицей шрифта),
    # тут физически не может быть проблемы с "нечитаемой" кодировкой шрифта.
    # Понятия "страница" в формате нет — весь текст идёт одним куском.
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(parts)
    if not text.strip():
        return []
    return [{"page_number": None, "text": text}]
